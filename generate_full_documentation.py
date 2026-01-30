"""
Скрипт для генерации полной документации проекта с вариантами изделия
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Установить границы ячейки таблицы"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def format_table(table):
    """Форматировать таблицу: черные границы, Times New Roman 14, запрет разрыва"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    border_settings = {
        'top': {'sz': '12', 'val': 'single', 'color': '000000'},
        'left': {'sz': '12', 'val': 'single', 'color': '000000'},
        'bottom': {'sz': '12', 'val': 'single', 'color': '000000'},
        'right': {'sz': '12', 'val': 'single', 'color': '000000'}
    }

    for row in table.rows:
        tr = row._element
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)

        for cell in row.cells:
            set_cell_border(cell, **border_settings)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 0, 0)

def format_heading(heading, size=14):
    """Форматировать заголовок"""
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True

def add_formatted_paragraph(doc, text, bold=False):
    """Добавить форматированный параграф"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold:
        run.font.bold = True
    return p

def add_ascii_diagram(doc, diagram_text):
    """Добавить ASCII-диаграмму"""
    p = doc.add_paragraph(diagram_text)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    return p

def create_documentation():
    """Создать документ с документацией"""
    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0, 0, 0)

    # Заголовок документа
    title = doc.add_heading('Документация проекта: Telegram-бот для отслеживания тренировок', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_heading(title, 16)

    # 1. Интервью
    h1 = doc.add_heading('Интервью «Ваше мнение»', 1)
    format_heading(h1)

    add_formatted_paragraph(doc, 'Я провел опрос спортсменов и тренера по бегу, плаванию и триатлону с целью выяснить отношение к будущему продукту. Результаты в Таблице 1 помогли точнее определить функциональность бота и приоритеты разработки.')

    # Таблица 1
    h2 = doc.add_heading('Таблица № 1. Результаты опроса целевой аудитории', 2)
    format_heading(h2)

    table1 = doc.add_table(rows=4, cols=3)
    table1.autofit = True

    headers1 = table1.rows[0].cells
    headers1[0].text = 'Опрашиваемые люди'
    headers1[1].text = 'Хотели бы вы использовать этот бот для себя? Почему?'
    headers1[2].text = 'Что бы вы хотели изменить или добавить?'

    for cell in headers1:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    table1.rows[1].cells[0].text = 'Спортсмены-любители'
    table1.rows[1].cells[1].text = 'Да. Это необходимый инструмент для самоконтроля и анализа своих тренировок. Удобно вести дневник тренировок прямо в Telegram, не переключаясь между приложениями.'
    table1.rows[1].cells[2].text = 'Необходимо добавить автоматический расчет пульсовых зон и напоминания о регулярном внесении показателей здоровья (пульс, вес, сон). Важна возможность экспорта данных в PDF для отправки тренеру.'

    table1.rows[2].cells[0].text = 'Тренер'
    table1.rows[2].cells[1].text = 'Как для тренера это необходимая вещь для контроля тренировочного процесса спортсменов и планирования нагрузок. После каждой тренировки можно провести анализ показателей и скорректировать план подготовки. Важна функция добавления тренировок и комментариев к ним для учеников.'
    table1.rows[2].cells[2].text = 'Необходимо оснастить бота функцией просмотра статистики ученика за разные периоды (неделя, месяц, год) с графиками объема и темпа. Также нужна система уведомлений о предстоящих стартах учеников и возможность отслеживать их прогресс к целевым результатам.'

    table1.rows[3].cells[0].text = 'Спортсмены-разрядники'
    table1.rows[3].cells[1].text = 'Да, особенно важна возможность отслеживания выполнения разрядных нормативов ЕВСК и регистрации на официальные соревнования через бота. Это экономит время на поиске стартов и расчете квалификации.'
    table1.rows[3].cells[2].text = 'Добавить интеграцию с большим количеством сервисов регистрации на соревнования (RussiaRunning, Timerman, HeroLeague). Важна возможность установки целевого времени и отслеживания прогресса к нему.'

    format_table(table1)

    # Таблица 2
    h2_2 = doc.add_heading('Таблица № 2. Анализ достоинств и недостатков', 2)
    format_heading(h2_2)

    table2 = doc.add_table(rows=6, cols=3)
    table2.autofit = True

    headers2 = table2.rows[0].cells
    headers2[0].text = 'Достоинства'
    headers2[1].text = 'Недостатки'
    headers2[2].text = 'Устранение недостатков'

    for cell in headers2:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    data2 = [
        ['Интеграция в привычный мессенджер (Telegram), не требует установки отдельного приложения',
         'Требует постоянного интернет-соединения для работы',
         'Добавить возможность офлайн-кэширования данных и отложенной синхронизации'],
        ['Многофункциональность: тренировочный дневник, здоровье, соревнования, связь с тренером',
         'Сложность навигации при большом количестве функций',
         'Оптимизация структуры меню, добавление быстрых команд и inline-кнопок'],
        ['Автоматические напоминания и уведомления о важных событиях',
         'Уведомления могут быть избыточными для некоторых пользователей',
         'Гибкая настройка уведомлений в разделе настроек'],
        ['Поддержка нескольких видов спорта (бег, плавание, велоспорт, триатлон, силовые)',
         'Специфичные поля для каждого вида спорта могут усложнить интерфейс',
         'Адаптивные формы ввода в зависимости от выбранного вида активности'],
        ['Система рейтингов и достижений для мотивации',
         'Не все пользователи заинтересованы в соревновательном аспекте',
         'Возможность отключения рейтингов в настройках']
    ]

    for i, row_data in enumerate(data2, start=1):
        for j, text in enumerate(row_data):
            table2.rows[i].cells[j].text = text

    format_table(table2)

    # НОВЫЙ РАЗДЕЛ: Возможные варианты изделия
    h1_variants = doc.add_heading('Возможные варианты изделия', 1)
    format_heading(h1_variants)

    add_formatted_paragraph(doc, 'Рассмотрены три варианта технической реализации системы, отличающиеся архитектурным подходом, масштабируемостью и сложностью развертывания.')

    # Вариант 1: Монолитная архитектура
    h2_v1 = doc.add_heading('Вариант 1: Монолитная архитектура с локальной базой данных', 2)
    format_heading(h2_v1)

    diagram1 = """
    ┌─────────────────────────────────────────────────────────┐
    │              Telegram Bot Application                   │
    │  ┌───────────────────────────────────────────────────┐  │
    │  │                  main.py                          │  │
    │  │          (инициализация, запуск)                  │  │
    │  └───────────────┬───────────────────────────────────┘  │
    │                  │                                       │
    │  ┌───────────────┼───────────────────────────────────┐  │
    │  │               │                                   │  │
    │  │    ┌──────────▼──────────┐    ┌─────────────────▼┐  │
    │  │    │   Handlers          │    │   Database       │  │
    │  │    │   (bot/, coach/,    │◄───┤   (SQLite)       │  │
    │  │    │   competitions/)    │    │   - queries.py   │  │
    │  │    └──────────┬──────────┘    │   - models.py    │  │
    │  │               │                └──────────────────┘  │
    │  │    ┌──────────▼──────────┐                          │
    │  │    │   FSM States        │                          │
    │  │    │   (bot/fsm.py)      │                          │
    │  │    └──────────┬──────────┘                          │
    │  │               │                                      │
    │  │    ┌──────────▼──────────┐                          │
    │  │    │   Schedulers        │                          │
    │  │    │   (notifications/,  │                          │
    │  │    │   ratings/)         │                          │
    │  │    └─────────────────────┘                          │
    │  └──────────────────────────────────────────────────────┘
    │                          │                               │
    └──────────────────────────┼───────────────────────────────┘
                               │
                    ┌──────────▼──────────┐
                    │   Telegram API      │
                    └─────────────────────┘
    """

    add_ascii_diagram(doc, diagram1)

    p = add_formatted_paragraph(doc, 'Достоинства:', bold=True)
    doc.add_paragraph('Простота развертывания - один файл приложения', style='List Bullet')
    doc.add_paragraph('Не требует отдельного сервера базы данных', style='List Bullet')
    doc.add_paragraph('Быстрая разработка и отладка', style='List Bullet')
    doc.add_paragraph('Низкие требования к ресурсам', style='List Bullet')

    for paragraph in doc.paragraphs[-4:]:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    p = add_formatted_paragraph(doc, 'Недостатки:', bold=True)
    doc.add_paragraph('Ограниченная масштабируемость при росте пользователей', style='List Bullet')
    doc.add_paragraph('SQLite не оптимален для высоких нагрузок', style='List Bullet')
    doc.add_paragraph('Сложность горизонтального масштабирования', style='List Bullet')

    for paragraph in doc.paragraphs[-3:]:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Вариант 2: Клиент-сервер с внешней БД
    h2_v2 = doc.add_heading('Вариант 2: Клиент-серверная архитектура с PostgreSQL', 2)
    format_heading(h2_v2)

    diagram2 = """
    ┌─────────────────────────┐         ┌─────────────────────────┐
    │   Telegram Bot          │         │   Database Server       │
    │                         │         │                         │
    │  ┌──────────────────┐   │         │  ┌──────────────────┐   │
    │  │  main.py         │   │         │  │   PostgreSQL     │   │
    │  │  (aiogram 3.x)   │   │         │  │                  │   │
    │  └────────┬─────────┘   │         │  │  - users         │   │
    │           │              │         │  │  - trainings     │   │
    │  ┌────────▼─────────┐   │         │  │  - competitions  │   │
    │  │  Handlers        │   │◄────────┤  │  - health        │   │
    │  │  - bot/          │   │  async  │  │  - ratings       │   │
    │  │  - coach/        │   │  queries│  └──────────────────┘   │
    │  │  - competitions/ │   │         │                         │
    │  └────────┬─────────┘   │         │  ┌──────────────────┐   │
    │           │              │         │  │  Backups         │   │
    │  ┌────────▼─────────┐   │         │  │  (автоматические)│   │
    │  │  Database Queries│───┼─────────┤  └──────────────────┘   │
    │  │  (asyncpg)       │   │         │                         │
    │  └──────────────────┘   │         └─────────────────────────┘
    │                         │
    │  ┌──────────────────┐   │
    │  │  Schedulers      │   │
    │  │  (background)    │   │
    │  └──────────────────┘   │
    └─────────────────────────┘
                │
    ┌───────────▼───────────┐
    │   Telegram API        │
    └───────────────────────┘
    """

    add_ascii_diagram(doc, diagram2)

    p = add_formatted_paragraph(doc, 'Достоинства:', bold=True)
    doc.add_paragraph('Высокая производительность PostgreSQL для больших объемов данных', style='List Bullet')
    doc.add_paragraph('Поддержка репликации и резервного копирования', style='List Bullet')
    doc.add_paragraph('Лучшая надежность и целостность данных', style='List Bullet')
    doc.add_paragraph('Возможность использования продвинутых SQL-функций', style='List Bullet')

    for paragraph in doc.paragraphs[-4:]:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    p = add_formatted_paragraph(doc, 'Недостатки:', bold=True)
    doc.add_paragraph('Требует отдельный сервер для базы данных', style='List Bullet')
    doc.add_paragraph('Усложнение развертывания и настройки', style='List Bullet')
    doc.add_paragraph('Дополнительные расходы на хостинг БД', style='List Bullet')
    doc.add_paragraph('Необходимость администрирования PostgreSQL', style='List Bullet')

    for paragraph in doc.paragraphs[-4:]:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Вариант 3: Микросервисная архитектура
    h2_v3 = doc.add_heading('Вариант 3: Микросервисная архитектура', 2)
    format_heading(h2_v3)

    diagram3 = """
    ┌─────────────────┐      ┌──────────────────┐      ┌─────────────────┐
    │  Bot Service    │      │  API Gateway     │      │  Database       │
    │                 │      │  (FastAPI)       │      │  (PostgreSQL)   │
    │  - Telegram API │◄─────┤                  │      │                 │
    │  - Handlers     │      │  - Auth          │      │  - Users        │
    │  - FSM          │      │  - Routing       │      │  - Trainings    │
    └────────┬────────┘      └────────┬─────────┘      │  - Health       │
             │                        │                 └────────▲────────┘
             │                        │                          │
             │               ┌────────▼─────────┐                │
             └───────────────┤  Training Service│────────────────┘
                             │                  │
                             │  - Add training  │
                             │  - Statistics    │
                             │  - PDF export    │
                             └────────┬─────────┘
                                      │
                             ┌────────▼─────────┐
                             │ Competition      │
                             │ Service          │────────────────┐
                             │                  │                │
                             │ - Search         │                │
                             │ - Registration   │                │
                             └────────┬─────────┘                │
                                      │                          │
                             ┌────────▼─────────┐      ┌─────────▼────────┐
                             │ Notification     │      │  External APIs   │
                             │ Service          │      │                  │
                             │                  │      │ - RussiaRunning  │
                             │ - Schedulers     │      │ - Timerman       │
                             │ - Reminders      │      │ - HeroLeague     │
                             └──────────────────┘      └──────────────────┘
    """

    add_ascii_diagram(doc, diagram3)

    p = add_formatted_paragraph(doc, 'Достоинства:', bold=True)
    doc.add_paragraph('Независимое масштабирование каждого сервиса', style='List Bullet')
    doc.add_paragraph('Высокая отказоустойчивость - падение одного сервиса не влияет на другие', style='List Bullet')
    doc.add_paragraph('Возможность использования разных технологий для разных сервисов', style='List Bullet')
    doc.add_paragraph('Упрощение командной разработки', style='List Bullet')

    for paragraph in doc.paragraphs[-4:]:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    p = add_formatted_paragraph(doc, 'Недостатки:', bold=True)
    doc.add_paragraph('Высокая сложность развертывания и поддержки', style='List Bullet')
    doc.add_paragraph('Требует оркестрацию (Docker, Kubernetes)', style='List Bullet')
    doc.add_paragraph('Повышенные требования к инфраструктуре', style='List Bullet')
    doc.add_paragraph('Избыточность для малого количества пользователей', style='List Bullet')

    for paragraph in doc.paragraphs[-4:]:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Таблица сравнения вариантов изделия
    h2_comparison = doc.add_heading('Таблица № 3. Сравнение вариантов изделия', 2)
    format_heading(h2_comparison)

    table3 = doc.add_table(rows=4, cols=4)
    table3.autofit = True

    headers3 = table3.rows[0].cells
    headers3[0].text = '№ варианта'
    headers3[1].text = 'Достоинства'
    headers3[2].text = 'Недостатки'
    headers3[3].text = 'Обоснование выбора'

    for cell in headers3:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    # Вариант 1
    table3.rows[1].cells[0].text = '1. Монолитная архитектура'
    table3.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    table3.rows[1].cells[1].text = 'Простота развертывания, низкие требования к ресурсам, быстрая разработка'
    table3.rows[1].cells[2].text = 'Ограниченная масштабируемость, SQLite не оптимален для высоких нагрузок'
    table3.rows[1].cells[3].text = 'ВЫБРАН ДЛЯ РЕАЛИЗАЦИИ. Оптимален для начального этапа и среднего количества пользователей (до 10000). Позволяет быстро запустить проект и протестировать функционал. При необходимости возможна миграция на Вариант 2.'
    table3.rows[1].cells[3].paragraphs[0].runs[0].font.bold = True

    # Вариант 2
    table3.rows[2].cells[0].text = '2. Клиент-сервер'
    table3.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True
    table3.rows[2].cells[1].text = 'Высокая производительность, надежность, поддержка репликации'
    table3.rows[2].cells[2].text = 'Требует отдельный сервер БД, усложнение развертывания, дополнительные расходы'
    table3.rows[2].cells[3].text = 'Рекомендуется при росте до 10000+ пользователей. Обеспечивает лучшую производительность и надежность.'

    # Вариант 3
    table3.rows[3].cells[0].text = '3. Микросервисная'
    table3.rows[3].cells[0].paragraphs[0].runs[0].font.bold = True
    table3.rows[3].cells[1].text = 'Независимое масштабирование, отказоустойчивость, гибкость'
    table3.rows[3].cells[2].text = 'Высокая сложность, требует оркестрацию, избыточность для малых проектов'
    table3.rows[3].cells[3].text = 'Подходит для масштабных проектов с 50000+ пользователей и распределенной командой разработки.'

    format_table(table3)

    # Результаты исследования
    h1_2 = doc.add_heading('Результаты исследования и учет пожеланий потребителя', 1)
    format_heading(h1_2)

    add_formatted_paragraph(doc, 'В результате исследования проекта и учета пожеланий потребителя при конструировании следует учесть следующие факторы:')

    h2_3 = doc.add_heading('Схема 1. Факторы проектирования', 2)
    format_heading(h2_3)

    scheme = """                    ┌──────────────────────────────────┐
                    │   Требования к боту              │
                    └──────────────┬───────────────────┘
                                   │
                ┌──────────────────┼──────────────────┐
                │                  │                  │
       ┌────────▼────────┐ ┌──────▼──────┐  ┌────────▼────────┐
       │  Функциональность│ │  Удобство   │  │   Технические   │
       │                  │ │ использования│  │   требования    │
       └────────┬────────┘ └──────┬──────┘  └────────┬────────┘
                │                  │                  │
        ┌───────┴───────┐  ┌──────┴──────┐   ┌──────┴────────┐
        │               │  │             │   │               │
   ┌────▼────┐   ┌─────▼──┐  ┌──────▼────┐ ┌─▼────┐  ┌──────▼─────┐
   │Дневник  │   │Здоровье│  │Интуитивная│ │Async/│  │Масштабиру- │
   │тренировок│   │        │  │навигация  │ │await │  │емость БД   │
   └─────────┘   └────────┘  └───────────┘ └──────┘  └────────────┘

   ┌─────────┐   ┌────────┐  ┌───────────┐ ┌──────┐  ┌────────────┐
   │Соревно- │   │Связь с │  │Быстрый    │ │FSM   │  │Уведомления │
   │вания    │   │тренером│  │отклик     │ │States│  │scheduler   │
   └─────────┘   └────────┘  └───────────┘ └──────┘  └────────────┘

   ┌─────────┐   ┌────────┐  ┌───────────┐ ┌──────┐  ┌────────────┐
   │Рейтинги │   │Настройки│  │Минимум    │ │aiogram│ │Резервное   │
   │         │   │         │  │шагов      │ │3.x   │  │копирование │
   └─────────┘   └────────┘  └───────────┘ └──────┘  └────────────┘"""

    add_ascii_diagram(doc, scheme)

    # Дополнительные факторы
    h1_4 = doc.add_heading('Дополнительные факторы, учтенные при разработке', 1)
    format_heading(h1_4)

    factors = [
        ('Система единиц измерения', [
            'Поддержка метрической и имперской системы',
            'Автоматический пересчет темпа и дистанций',
            'Настройка в профиле пользователя'
        ]),
        ('Локализация и часовые пояса', [
            'Поддержка русского языка',
            'Учет часового пояса пользователя для уведомлений',
            'Форматирование дат по местным стандартам'
        ]),
        ('Интеграция с внешними API', [
            'RussiaRunning API для российских соревнований',
            'Timerman API для дополнительных стартов',
            'HeroLeague API для триатлона и мультиспорта',
            'RegPlace для региональных соревнований'
        ]),
        ('Система квалификации', [
            'Автоматический расчет разрядных нормативов ЕВСК',
            'Отслеживание прогресса к целевому разряду',
            'Уведомления о выполнении норматива'
        ]),
        ('Безопасность и конфиденциальность', [
            'Все данные хранятся локально в SQLite',
            'Связь тренер-ученик через уникальные коды',
            'Возможность удаления всех данных'
        ])
    ]

    for i, (title, items) in enumerate(factors, 1):
        h2_f = doc.add_heading(f'{i}. {title}', 2)
        format_heading(h2_f)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
            for run in doc.paragraphs[-1].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Итоговое решение
    h1_5 = doc.add_heading('Итоговое решение', 1)
    format_heading(h1_5)

    p = doc.add_paragraph()
    run1 = p.add_run('На основе проведенного исследования и анализа был выбран ')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(14)
    run1.font.color.rgb = RGBColor(0, 0, 0)

    run2 = p.add_run('Вариант 1 (Монолитная архитектура)')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0, 0, 0)
    run2.font.bold = True

    run3 = p.add_run(' как основная архитектура системы с учетом следующих факторов:')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0, 0, 0)

    reasons = [
        'Оптимальное соотношение простоты разработки и функциональности',
        'Низкие требования к инфраструктуре на начальном этапе',
        'Быстрое развертывание и тестирование',
        'Возможность миграции на клиент-серверную архитектуру при необходимости',
        'Достаточная производительность для целевой аудитории'
    ]

    for reason in reasons:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(reason)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

    add_formatted_paragraph(doc, '\nДля обеспечения будущего роста предусмотрена модульная структура кода, позволяющая легко мигрировать на Вариант 2 (PostgreSQL) при достижении 10000 пользователей.')

    # Информация об авторе
    doc.add_paragraph()
    add_formatted_paragraph(doc, '─' * 50)

    p = doc.add_paragraph()
    r1 = p.add_run('Автор: ')
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(14)
    r1.font.color.rgb = RGBColor(0, 0, 0)
    r1.font.bold = True
    r2 = p.add_run('Ученик')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    r1 = p.add_run('Дата создания: ')
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(14)
    r1.font.color.rgb = RGBColor(0, 0, 0)
    r1.font.bold = True
    r2 = p.add_run('2026-01-29')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    r1 = p.add_run('Статус проекта: ')
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(14)
    r1.font.color.rgb = RGBColor(0, 0, 0)
    r1.font.bold = True
    r2 = p.add_run('Training assistant реализован')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0, 0, 0)

    # Сохранение
    doc.save('D:\\desktop\\coding\\trainingdairy_bot-4\\Документация_полная.docx')
    print('Документ успешно создан: Документация_полная.docx')

if __name__ == '__main__':
    create_documentation()
