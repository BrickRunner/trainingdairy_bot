"""
Скрипт для генерации документации проекта в формате .docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Установить границы ячейки таблицы"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def format_table(table):
    """Форматировать таблицу: черные границы, Times New Roman 14, запрет разрыва"""
    # Запрет разрыва таблицы между страницами
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Границы
    border_settings = {
        'top': {'sz': '12', 'val': 'single', 'color': '000000'},
        'left': {'sz': '12', 'val': 'single', 'color': '000000'},
        'bottom': {'sz': '12', 'val': 'single', 'color': '000000'},
        'right': {'sz': '12', 'val': 'single', 'color': '000000'}
    }

    for row in table.rows:
        # Запрет разрыва строк
        tr = row._element
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)

        for cell in row.cells:
            set_cell_border(cell, **border_settings)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 0, 0)

def format_heading(heading, size=14):
    """Форматировать заголовок"""
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True

def add_formatted_paragraph(doc, text, bold=False):
    """Добавить форматированный параграф"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold:
        run.font.bold = True
    return p

def create_documentation():
    """Создать документ с документацией"""
    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0, 0, 0)

    # Заголовок документа
    title = doc.add_heading('Документация проекта: Telegram-бот для отслеживания тренировок', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_heading(title, 16)

    # 1. Интервью
    h1 = doc.add_heading('Интервью «Ваше мнение»', 1)
    format_heading(h1)

    add_formatted_paragraph(doc, 'Я провел опрос спортсменов и тренера по бегу, плаванию и триатлону с целью выяснить отношение к будущему продукту. Результаты в Таблице 1 помогли точнее определить функциональность бота и приоритеты разработки.')

    # Таблица 1
    h2 = doc.add_heading('Таблица № 1. Результаты опроса целевой аудитории', 2)
    format_heading(h2)

    table1 = doc.add_table(rows=4, cols=3)
    table1.autofit = True

    headers1 = table1.rows[0].cells
    headers1[0].text = 'Опрашиваемые люди'
    headers1[1].text = 'Хотели бы вы использовать этот бот для себя? Почему?'
    headers1[2].text = 'Что бы вы хотели изменить или добавить?'

    for cell in headers1:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    # Данные таблицы 1
    table1.rows[1].cells[0].text = 'Спортсмены-любители'
    table1.rows[1].cells[1].text = 'Да. Это необходимый инструмент для самоконтроля и анализа своих тренировок. Удобно вести дневник тренировок прямо в Telegram, не переключаясь между приложениями.'
    table1.rows[1].cells[2].text = 'Необходимо добавить автоматический расчет пульсовых зон и напоминания о регулярном внесении показателей здоровья (пульс, вес, сон). Важна возможность экспорта данных в PDF для отправки тренеру.'

    table1.rows[2].cells[0].text = 'Тренер'
    table1.rows[2].cells[1].text = 'Как для тренера это необходимая вещь для контроля тренировочного процесса спортсменов и планирования нагрузок. После каждой тренировки можно провести анализ показателей и скорректировать план подготовки. Важна функция добавления тренировок и комментариев к ним для учеников.'
    table1.rows[2].cells[2].text = 'Необходимо оснастить бота функцией просмотра статистики ученика за разные периоды (неделя, месяц, год) с графиками объема и темпа. Также нужна система уведомлений о предстоящих стартах учеников и возможность отслеживать их прогресс к целевым результатам.'

    table1.rows[3].cells[0].text = 'Спортсмены-разрядники'
    table1.rows[3].cells[1].text = 'Да, особенно важна возможность отслеживания выполнения разрядных нормативов ЕВСК и регистрации на официальные соревнования через бота. Это экономит время на поиске стартов и расчете квалификации.'
    table1.rows[3].cells[2].text = 'Добавить интеграцию с большим количеством сервисов регистрации на соревнования (RussiaRunning, Timerman, HeroLeague). Важна возможность установки целевого времени и отслеживания прогресса к нему.'

    format_table(table1)

    # Таблица 2
    h2_2 = doc.add_heading('Таблица № 2. Анализ достоинств и недостатков', 2)
    format_heading(h2_2)

    table2 = doc.add_table(rows=6, cols=3)
    table2.autofit = True

    headers2 = table2.rows[0].cells
    headers2[0].text = 'Достоинства'
    headers2[1].text = 'Недостатки'
    headers2[2].text = 'Устранение недостатков'

    for cell in headers2:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    data2 = [
        ['Интеграция в привычный мессенджер (Telegram), не требует установки отдельного приложения',
         'Требует постоянного интернет-соединения для работы',
         'Добавить возможность офлайн-кэширования данных и отложенной синхронизации'],
        ['Многофункциональность: тренировочный дневник, здоровье, соревнования, связь с тренером',
         'Сложность навигации при большом количестве функций',
         'Оптимизация структуры меню, добавление быстрых команд и inline-кнопок'],
        ['Автоматические напоминания и уведомления о важных событиях',
         'Уведомления могут быть избыточными для некоторых пользователей',
         'Гибкая настройка уведомлений в разделе настроек'],
        ['Поддержка нескольких видов спорта (бег, плавание, велоспорт, триатлон, силовые)',
         'Специфичные поля для каждого вида спорта могут усложнить интерфейс',
         'Адаптивные формы ввода в зависимости от выбранного вида активности'],
        ['Система рейтингов и достижений для мотивации',
         'Не все пользователи заинтересованы в соревновательном аспекте',
         'Возможность отключения рейтингов в настройках']
    ]

    for i, row_data in enumerate(data2, start=1):
        for j, text in enumerate(row_data):
            table2.rows[i].cells[j].text = text

    format_table(table2)

    # Результаты исследования
    h1_2 = doc.add_heading('Результаты исследования и учет пожеланий потребителя', 1)
    format_heading(h1_2)

    add_formatted_paragraph(doc, 'В результате исследования проекта и учета пожеланий потребителя при конструировании следует учесть следующие факторы:')

    h2_3 = doc.add_heading('Схема 1. Факторы проектирования', 2)
    format_heading(h2_3)

    scheme = """                    ┌──────────────────────────────────┐
                    │   Требования к боту              │
                    └──────────────┬───────────────────┘
                                   │
                ┌──────────────────┼──────────────────┐
                │                  │                  │
       ┌────────▼────────┐ ┌──────▼──────┐  ┌────────▼────────┐
       │  Функциональность│ │  Удобство   │  │   Технические   │
       │                  │ │ использования│  │   требования    │
       └────────┬────────┘ └──────┬──────┘  └────────┬────────┘
                │                  │                  │
        ┌───────┴───────┐  ┌──────┴──────┐   ┌──────┴────────┐
        │               │  │             │   │               │
   ┌────▼────┐   ┌─────▼──┐  ┌──────▼────┐ ┌─▼────┐  ┌──────▼─────┐
   │Дневник  │   │Здоровье│  │Интуитивная│ │Async/│  │Масштабиру- │
   │тренировок│   │        │  │навигация  │ │await │  │емость БД   │
   └─────────┘   └────────┘  └───────────┘ └──────┘  └────────────┘

   ┌─────────┐   ┌────────┐  ┌───────────┐ ┌──────┐  ┌────────────┐
   │Соревно- │   │Связь с │  │Быстрый    │ │FSM   │  │Уведомления │
   │вания    │   │тренером│  │отклик     │ │States│  │scheduler   │
   └─────────┘   └────────┘  └───────────┘ └──────┘  └────────────┘

   ┌─────────┐   ┌────────┐  ┌───────────┐ ┌──────┐  ┌────────────┐
   │Рейтинги │   │Настройки│  │Минимум    │ │aiogram│ │Резервное   │
   │         │   │         │  │шагов      │ │3.x   │  │копирование │
   └─────────┘   └────────┘  └───────────┘ └──────┘  └────────────┘"""

    p = doc.add_paragraph(scheme)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # Варианты реализации
    h1_3 = doc.add_heading('Возможные варианты реализации интерфейса', 1)
    format_heading(h1_3)

    h2_4 = doc.add_heading('Таблица 3. Варианты структуры меню', 2)
    format_heading(h2_4)

    # Вариант 1
    h3_1 = doc.add_heading('Вариант 1: Иерархическое меню с разделами', 3)
    format_heading(h3_1)

    menu1 = """Главное меню
├── Тренировки
│   ├── Добавить тренировку
│   ├── Статистика
│   ├── Графики
│   └── Экспорт в PDF
├── Соревнования
│   ├── Поиск соревнований
│   ├── Мои регистрации
│   ├── Результаты
│   └── Добавить свое
├── Здоровье
│   ├── Добавить показатели
│   ├── История
│   └── Экспорт
├── Тренер
│   ├── Связь с тренером
│   ├── Мои ученики
│   └── Комментарии
├── Рейтинги
└── Настройки"""

    p = doc.add_paragraph(menu1)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    p = add_formatted_paragraph(doc, 'Достоинства:', bold=True)
    doc.add_paragraph('Логичная структура с четким разделением функционала', style='List Bullet')
    doc.add_paragraph('Легко найти нужный раздел', style='List Bullet')
    doc.add_paragraph('Возможность расширения функционала', style='List Bullet')

    # Форматируем списки
    for paragraph in doc.paragraphs[-3:]:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    p = add_formatted_paragraph(doc, 'Недостатки:', bold=True)
    doc.add_paragraph('Требуется несколько нажатий для доступа к функциям', style='List Bullet')
    doc.add_paragraph('Может быть избыточным для простых задач', style='List Bullet')

    for paragraph in doc.paragraphs[-2:]:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Вариант 2
    h3_2 = doc.add_heading('Вариант 2: Плоское меню с быстрым доступом', 3)
    format_heading(h3_2)

    menu2 = """Главное меню
├── Добавить тренировку
├── Найти соревнование
├── Здоровье
├── Моя статистика
├── Связь с тренером
├── Рейтинги
└── Все функции"""

    p = doc.add_paragraph(menu2)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    p = add_formatted_paragraph(doc, 'Достоинства:', bold=True)
    doc.add_paragraph('Быстрый доступ к основным функциям', style='List Bullet')
    doc.add_paragraph('Меньше кликов для частых действий', style='List Bullet')
    doc.add_paragraph('Простота использования', style='List Bullet')

    for paragraph in doc.paragraphs[-3:]:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    p = add_formatted_paragraph(doc, 'Недостатки:', bold=True)
    doc.add_paragraph('Ограниченное количество функций в главном меню', style='List Bullet')
    doc.add_paragraph('Дополнительные функции скрыты в подменю', style='List Bullet')

    for paragraph in doc.paragraphs[-2:]:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Таблица 4
    h2_5 = doc.add_heading('Таблица 4. Сравнение вариантов реализации интерфейса', 2)
    format_heading(h2_5)

    table4 = doc.add_table(rows=3, cols=4)
    table4.autofit = True

    headers4 = table4.rows[0].cells
    headers4[0].text = '№ варианта'
    headers4[1].text = 'Достоинства'
    headers4[2].text = 'Недостатки'
    headers4[3].text = 'Обоснование выбора лучшего варианта'

    for cell in headers4:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    # Вариант 1
    table4.rows[1].cells[0].text = '1. Иерархическое меню'
    table4.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    table4.rows[1].cells[1].text = 'Возможность отслеживать весь функционал, логичная структура, масштабируемость'
    table4.rows[1].cells[2].text = 'Больше шагов до нужной функции, может показаться перегруженным новым пользователям'
    table4.rows[1].cells[3].text = 'ВЫБРАН КАК ОСНОВНОЙ. Удобство структурированной навигации перевешивает недостаток дополнительных кликов. Пользователи быстро запоминают структуру. Возможность добавления inline-кнопок компенсирует количество шагов.'
    table4.rows[1].cells[3].paragraphs[0].runs[0].font.bold = True

    # Вариант 2
    table4.rows[2].cells[0].text = '2. Плоское меню'
    table4.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True
    table4.rows[2].cells[1].text = 'Быстрый доступ к основным функциям, минимум кликов'
    table4.rows[2].cells[2].text = 'Ограниченное количество функций в главном меню, сложность расширения функционала'
    table4.rows[2].cells[3].text = 'Подходит для упрощенной версии, но ограничивает развитие проекта. Используется частично через команды быстрого доступа (/add, /stats).'

    format_table(table4)

    # Дополнительные факторы
    h1_4 = doc.add_heading('Дополнительные факторы, учтенные при разработке', 1)
    format_heading(h1_4)

    factors = [
        ('Система единиц измерения', [
            'Поддержка метрической и имперской системы',
            'Автоматический пересчет темпа и дистанций',
            'Настройка в профиле пользователя'
        ]),
        ('Локализация и часовые пояса', [
            'Поддержка русского языка',
            'Учет часового пояса пользователя для уведомлений',
            'Форматирование дат по местным стандартам'
        ]),
        ('Интеграция с внешними API', [
            'RussiaRunning API для российских соревнований',
            'Timerman API для дополнительных стартов',
            'HeroLeague API для триатлона и мультиспорта',
            'RegPlace для региональных соревнований'
        ]),
        ('Система квалификации', [
            'Автоматический расчет разрядных нормативов ЕВСК',
            'Отслеживание прогресса к целевому разряду',
            'Уведомления о выполнении норматива'
        ]),
        ('Безопасность и конфиденциальность', [
            'Все данные хранятся локально в SQLite',
            'Связь тренер-ученик через уникальные коды',
            'Возможность удаления всех данных'
        ])
    ]

    for i, (title, items) in enumerate(factors, 1):
        h2_f = doc.add_heading(f'{i}. {title}', 2)
        format_heading(h2_f)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
            for run in doc.paragraphs[-1].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Итоговое решение
    h1_5 = doc.add_heading('Итоговое решение', 1)
    format_heading(h1_5)

    p = doc.add_paragraph()
    run1 = p.add_run('На основе проведенного исследования и анализа был выбран ')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(14)
    run1.font.color.rgb = RGBColor(0, 0, 0)

    run2 = p.add_run('Вариант 1 (Иерархическое меню)')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0, 0, 0)
    run2.font.bold = True

    run3 = p.add_run(' с добавлением следующих оптимизаций:')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0, 0, 0)

    optimizations = [
        ('Inline-кнопки', 'для быстрого доступа к частым действиям'),
        ('Быстрые команды', '(/add, /stats, /health) для опытных пользователей'),
        ('Контекстная навигация', 'возврат в предыдущий раздел без возвращения в главное меню'),
        ('Адаптивные формы', 'показ только релевантных полей в зависимости от типа активности'),
        ('FSM (Finite State Machine)', 'для управления многошаговыми процессами')
    ]

    for title, desc in optimizations:
        p = doc.add_paragraph(style='List Number')
        run_title = p.add_run(title)
        run_title.font.name = 'Times New Roman'
        run_title.font.size = Pt(14)
        run_title.font.color.rgb = RGBColor(0, 0, 0)
        run_title.font.bold = True

        run_desc = p.add_run(f' - {desc}')
        run_desc.font.name = 'Times New Roman'
        run_desc.font.size = Pt(14)
        run_desc.font.color.rgb = RGBColor(0, 0, 0)

    add_formatted_paragraph(doc, '\nЭти решения позволяют сочетать структурированность иерархического меню с удобством быстрого доступа, удовлетворяя потребности как новых, так и опытных пользователей.')

    # Информация об авторе
    doc.add_paragraph()
    add_formatted_paragraph(doc, '─' * 50)

    p = doc.add_paragraph()
    r1 = p.add_run('Автор: ')
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(14)
    r1.font.color.rgb = RGBColor(0, 0, 0)
    r1.font.bold = True
    r2 = p.add_run('Ученик')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    r1 = p.add_run('Дата создания: ')
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(14)
    r1.font.color.rgb = RGBColor(0, 0, 0)
    r1.font.bold = True
    r2 = p.add_run('2026-01-29')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    r1 = p.add_run('Статус проекта: ')
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(14)
    r1.font.color.rgb = RGBColor(0, 0, 0)
    r1.font.bold = True
    r2 = p.add_run('Training assistant реализован')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0, 0, 0)

    # Сохранение
    doc.save('D:\\desktop\\coding\\trainingdairy_bot-4\\Документация_проекта_v2.docx')
    print('Документ успешно создан: Документация_проекта_v2.docx')

if __name__ == '__main__':
    create_documentation()
