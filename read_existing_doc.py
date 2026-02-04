"""
Скрипт для чтения существующей документации
"""
from docx import Document
import sys

def read_docx(filename):
    """Читает и выводит содержимое DOCX файла"""
    try:
        doc = Document(filename)

        print(f"\n{'='*80}")
        print(f"Файл: {filename}")
        print(f"Количество параграфов: {len(doc.paragraphs)}")
        print(f"Количество таблиц: {len(doc.tables)}")
        print(f"{'='*80}\n")

        # Выводим содержимое
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:
                style = para.style.name
                print(f"[{i}] [{style}] {text[:200]}")

        # Информация о таблицах
        if doc.tables:
            print(f"\n{'='*80}")
            print("ТАБЛИЦЫ:")
            print(f"{'='*80}\n")
            for i, table in enumerate(doc.tables, 1):
                print(f"Таблица {i}: {len(table.rows)} строк × {len(table.columns)} столбцов")
                # Выводим первую строку (заголовки)
                if table.rows:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    print(f"  Заголовки: {headers}")

    except Exception as e:
        print(f"Ошибка при чтении файла: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    # Читаем "Документация к проекту (2).docx"
    read_docx('Документация к проекту (2).docx')
