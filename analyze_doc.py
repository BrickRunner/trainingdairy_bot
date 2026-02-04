"""
Анализ существующей документации
"""
from docx import Document

def analyze_docx(filename):
    """Анализирует структуру DOCX файла"""
    try:
        doc = Document(filename)

        output = []
        output.append("="*80)
        output.append(f"Файл: {filename}")
        output.append(f"Параграфов: {len(doc.paragraphs)}")
        output.append(f"Таблиц: {len(doc.tables)}")
        output.append("="*80)
        output.append("")

        # Анализируем структуру
        output.append("СТРУКТУРА ДОКУМЕНТА:")
        output.append("")

        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:
                style = para.style.name
                # Выводим только заголовки и важные элементы
                if 'Heading' in style or len(text) < 100:
                    indent = "  " * (int(style.replace('Heading ', '')) if 'Heading' in style else 0)
                    output.append(f"{i:3d}. {indent}[{style}] {text[:150]}")

        # Сохраняем в файл
        with open('doc_structure.txt', 'w', encoding='utf-8') as f:
            f.write('\n'.join(output))

        print("Структура сохранена в doc_structure.txt")

        # Также создаем полный дамп
        full_output = []
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:
                style = para.style.name
                full_output.append(f"[{i:3d}] [{style:20s}] {text}")

        with open('doc_full.txt', 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_output))

        print("Полный текст сохранен в doc_full.txt")

        # Информация о таблицах
        if doc.tables:
            table_info = ["\nТАБЛИЦЫ:"]
            for i, table in enumerate(doc.tables, 1):
                table_info.append(f"\nТаблица {i}: {len(table.rows)} строк × {len(table.columns)} столбцов")
                # Заголовки
                if table.rows:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    table_info.append(f"  Заголовки: {' | '.join(headers)}")

            with open('doc_tables.txt', 'w', encoding='utf-8') as f:
                f.write('\n'.join(table_info))
            print("Таблицы сохранены в doc_tables.txt")

    except Exception as e:
        print(f"Ошибка: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    analyze_docx('Документация к проекту (2).docx')
